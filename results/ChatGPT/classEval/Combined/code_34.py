from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class DocFileHandler:
    """
    A class to handle Word documents with functionalities for reading, writing, 
    modifying content, adding headings, and inserting tables.
    """

    def __init__(self, file_path):
        """
        Initialize DocFileHandler with the specified file path.
        :param file_path: str, path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Read the content of a Word document and return it as a string.
        :return: str, content of the Word document.
        """
        doc = Document(self.file_path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Write specified content to a Word document.
        :param content: str, text content to write.
        :param font_size: int, optional, font size of the text (default is 12).
        :param alignment: str, optional, text alignment ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if successful, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            p = doc.add_paragraph(content)
            run = p.runs[0]
            run.font.size = Pt(font_size)
            p.alignment = self._get_alignment_value(alignment)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing text: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Add a heading to the Word document.
        :param heading: str, text of the heading.
        :param level: int, optional, level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if successful, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Add a table to the Word document with specified data.
        :param data: list of lists, data to populate the table.
        :return: bool, True if successful, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=len(data), cols=len(data[0]))
            for row_idx, row in enumerate(data):
                for col_idx, cell_value in enumerate(row):
                    table.cell(row_idx, col_idx).text = cell_value
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Return the alignment value corresponding to the given alignment string.
        :param alignment: str, alignment string ('left', 'center', or 'right').
        :return: WD_PARAGRAPH_ALIGNMENT, alignment value.
        """
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)