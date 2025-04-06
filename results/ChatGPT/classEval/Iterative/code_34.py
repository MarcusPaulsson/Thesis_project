from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class DocFileHandler:
    """
    This class handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist.")
        self.file_path = file_path
        self.document = Document(file_path)

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        return '\n'.join(paragraph.text for paragraph in self.document.paragraphs)

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            paragraph = self.document.add_paragraph(content)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            paragraph.alignment = self._get_alignment_value(alignment)
            return True
        except Exception as e:
            print(f"Error writing text: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            self.document.add_heading(heading, level=level)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        if not data or not all(isinstance(row, list) for row in data):
            print("Invalid data format for table. Must be a list of lists.")
            return False

        try:
            table = self.document.add_table(rows=len(data), cols=len(data[0]))
            for row_idx, row in enumerate(data):
                for col_idx, cell_data in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(cell_data)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment_map = {
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT
        }
        return alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def save(self):
        """
        Saves the document to the specified file path.
        """
        try:
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False