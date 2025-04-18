from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        try:
            doc = Document(self.file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            return '\n'.join(full_text)
        except FileNotFoundError:
            return ""

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            doc = Document()
            paragraph = doc.add_paragraph(content)
            paragraph.alignment = self._get_alignment_value(alignment)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing to file: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level=level)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=1, cols=len(data[0]))
            for i, row_data in enumerate(data):
                for j, cell_data in enumerate(row_data):
                    table.cell(i, j).text = str(cell_data)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT