from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Union


class DocFileHandler:
    """
    Handles Word documents, providing functionalities for reading, writing,
    and modifying the content.
    """

    def __init__(self, file_path: str):
        """
        Initializes the DocFileHandler with the specified file path.

        Args:
            file_path: The path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self) -> str:
        """
        Reads the content of a Word document and returns it as a string.

        Returns:
            The content of the Word document, or an empty string if an error occurs.
        """
        try:
            document = Document(self.file_path)
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except FileNotFoundError:
            print(f"Error: File not found at path: {self.file_path}")
            return ""
        except Exception as e:
            print(f"Error reading document: {e}")
            return ""

    def write_text(self, content: str, font_size: int = 12, alignment: str = 'left') -> bool:
        """
        Writes the specified content to a Word document.  Overwrites existing content.

        Args:
            content: The text content to write.
            font_size: The font size of the text (default is 12).
            alignment: The alignment of the text ('left', 'center', or 'right'; default is 'left').

        Returns:
            True if the write operation is successful, False otherwise.
        """
        try:
            document = Document()
            paragraph = document.add_paragraph(content)
            paragraph.alignment = self._get_alignment_value(alignment)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing to document: {e}")
            return False

    def add_heading(self, heading: str, level: int = 1) -> bool:
        """
        Adds a heading to the Word document. Appends to existing content.

        Args:
            heading: The text of the heading.
            level: The level of the heading (1, 2, 3, etc.; default is 1).

        Returns:
            True if the heading is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path) if os.path.exists(self.file_path) else Document()
            document.add_heading(heading, level=level)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data: List[List[Union[str, int, float]]]) -> bool:
        """
        Adds a table to the Word document with the specified data. Appends to existing content.

        Args:
            data: A list of lists representing the table data.  Each inner list is a row.

        Returns:
            True if the table is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path) if os.path.exists(self.file_path) else Document()
            if not data:
                print("Warning: No data provided for the table.")
                document.save(self.file_path)
                return True

            table = document.add_table(rows=0, cols=len(data[0]))

            for row_data in data:
                row_cells = table.add_row().cells
                for i, item in enumerate(row_data):
                    row_cells[i].text = str(item)

            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment: str) -> int:
        """
        Returns the alignment value corresponding to the given alignment string.

        Args:
            alignment: The alignment string ('left', 'center', or 'right').

        Returns:
            The alignment value (WD_PARAGRAPH_ALIGNMENT).  Defaults to left if
            an invalid alignment string is provided.
        """
        alignment = alignment.lower()
        if alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT

import os
from docx import Document


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        try:
            document = Document(self.file_path)
            text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading document: {e}")
            return ""

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            document = Document()
            paragraph = document.add_paragraph(content)
            paragraph.alignment = self._get_alignment_value(alignment)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing to document: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path) if os.path.exists(self.file_path) else Document()
            document.add_heading(heading, level=level)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path) if os.path.exists(self.file_path) else Document()
            if not data:
                print("Warning: No data provided for the table.")
                document.save(self.file_path)
                return True

            table = document.add_table(rows=0, cols=len(data[0]))

            for row_data in data:
                row_cells = table.add_row().cells
                for i, item in enumerate(row_data):
                    row_cells[i].text = str(item)

            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment = alignment.lower()
        if alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT