from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


class DocFileHandler:
    """
    Handles Word documents, providing functionalities for reading, writing, and modifying content.
    """

    def __init__(self, file_path=None):
        """
        Initializes the DocFileHandler object with an optional file path.
        If no file_path is provided, a new document will be created.
        :param file_path: str, the path to the Word document file (optional).
        """
        self.file_path = file_path
        if file_path:
            try:
                self.document = Document(file_path)
            except Exception as e:
                print(f"Error opening document: {e}. Creating a new document.")
                self.document = Document()
        else:
            self.document = Document()

    def read_text(self):
        """
        Reads the content of the Word document and returns it as a string.
        :return: str, the content of the document, or None if an error occurs.
        """
        try:
            text = "\n".join(paragraph.text for paragraph in self.document.paragraphs)
            return text
        except Exception as e:
            print(f"Error reading document: {e}")
            return None

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to the Word document.  Creates a new paragraph.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.alignment = self._get_alignment_value(alignment)
            run = paragraph.add_run(content)
            run.font.size = Pt(font_size)
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing to document: {e}")
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            self.document.add_heading(heading, level=level)
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.  Must be rectangular.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            if not data or not data[0]:
                print("No data provided or empty table structure.")
                return False

            table = self.document.add_table(rows=0, cols=len(data[0]))

            for row_data in data:
                row_cells = table.add_row().cells
                if len(row_data) != len(data[0]):
                    print("Warning: Row with incorrect number of columns. Skipping.")
                    continue  # Skip rows with incorrect column counts

                for i, item in enumerate(row_data):
                    row_cells[i].text = str(item)

            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.  Defaults to LEFT.
        """
        alignment = alignment.lower()
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            print(f"Invalid alignment value: {alignment}. Defaulting to left.")
            return WD_PARAGRAPH_ALIGNMENT.LEFT  # Default to left alignment

    def save(self, file_path=None):
        """Saves the document.  If file_path is None, saves to the original file path.
           If no original file path was set, raises a ValueError.
        """
        if file_path:
            self.file_path = file_path
        if not self.file_path:
            raise ValueError("No file path specified for saving.")
        try:
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False

    def add_page_break(self):
        """Adds a page break to the document."""
        try:
            self.document.add_page_break()
            self.document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding page break: {e}")
            return False